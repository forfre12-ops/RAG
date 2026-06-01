"""HWP/DOCX 업로드 파싱 품질 테스트."""
import io
import json
import sys
import time
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

import httpx

API_URL = "http://localhost:8000"
API_KEY = "devkey"
ACTOR   = json.dumps({"tenant_id": "poc", "user_id": "test", "role": "system"})

results = []


def test_upload(filepath: Path, doc_type: str = "S1"):
    """파일 업로드 후 추출 품질 확인."""
    with httpx.Client(timeout=30.0) as cli:
        t0 = time.perf_counter()
        r = cli.post(
            f"{API_URL}/api/v1/documents",
            headers={"X-API-Key": API_KEY},
            data={"actor": ACTOR, "doc_type": doc_type},
            files={"file": (filepath.name, filepath.read_bytes(),
                            "application/octet-stream")},
        )
        elapsed_ms = (time.perf_counter() - t0) * 1000
        if r.status_code == 201:
            d = r.json()
            print(f"  [OK] {filepath.name}: "
                  f"method={d['extraction_method']} "
                  f"quality={d['extraction_quality']:.2f} "
                  f"chars={d['char_count']} "
                  f"{elapsed_ms:.0f}ms")
            return d
        else:
            print(f"  [FAIL] {filepath.name}: {r.status_code} - {r.text[:80]}")
            return None


# ── DOCX 생성 후 테스트 ───────────────────────────────────────
print("=== DOCX 업로드 테스트 ===")
try:
    from docx import Document
    import tempfile

    doc = Document()
    doc.add_heading("영업비밀 등급 S1 — 알고리즘 소스코드 보호 계획", level=1)
    doc.add_paragraph("본 문서는 1급 비밀로 분류된 핵심 알고리즘 소스코드 보호 계획서입니다.")
    doc.add_paragraph("공정 노하우 및 원가 구조 분석 결과가 포함되어 있습니다.")

    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "항목"
    table.cell(0, 1).text = "내용"
    table.cell(1, 0).text = "분류 등급"
    table.cell(1, 1).text = "S1 (1급 비밀)"
    table.cell(2, 0).text = "담당 부서"
    table.cell(2, 1).text = "기술보안팀"

    tmp = Path(tempfile.mktemp(suffix=".docx"))
    doc.save(tmp)
    result = test_upload(tmp, "S1")
    if result:
        results.append({"format": "docx", **result})
    tmp.unlink(missing_ok=True)
except ImportError:
    print("  python-docx 미설치")

# ── HWP 테스트 (파일 존재 시) ─────────────────────────────────
print("\n=== HWP 업로드 테스트 ===")
hwp_files = list(Path("datasets").rglob("*.hwp")) + list(Path("datasets").rglob("*.hwpx"))
if hwp_files:
    for f in hwp_files[:2]:
        result = test_upload(f, "S2")
        if result:
            results.append({"format": f.suffix, **result})
else:
    print("  HWP 파일 없음 — datasets/ 에 .hwp/.hwpx 파일 추가 시 테스트 가능")
    # rhwp 라이브러리 설치 여부 확인
    try:
        import rhwp  # noqa: F401
        print("  rhwp 라이브러리: 설치됨 (HWP 파서 준비 완료)")
    except ImportError:
        print("  rhwp 라이브러리: 미설치 (pip install rhwp 필요)")
    try:
        import hwp5  # noqa: F401
        print("  hwp5 라이브러리: 설치됨 (HWP5 폴백 준비 완료)")
    except ImportError:
        print("  hwp5 라이브러리: 미설치 (pip install pyhwp 필요)")

# ── 결과 저장 ─────────────────────────────────────────────────
Path("reports").mkdir(exist_ok=True)
Path("reports/hwp_docx_test.json").write_text(
    json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8")
print(f"\n완료: {len(results)}건 → reports/hwp_docx_test.json")
