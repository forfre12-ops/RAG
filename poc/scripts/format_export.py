"""골든셋 → 실제 문서 형식(DOCX/PDF) 변환 — 파이프라인(M2 추출) 견고성 검증용 서브셋.

분류 정확도는 텍스트로 측정하지만, End-to-End(파일 업로드→추출→분류)는 실제 형식이 필요.
DOCX(python-docx)·PDF(reportlab 한글 CID)는 생성 가능. HWP는 별도 방안(하단 주석).

실행:
  cd poc
  PYTHONPATH=src .venv/Scripts/python.exe scripts/format_export.py \
    --input datasets/gold/golden50.jsonl --limit 0
"""

from __future__ import annotations

import argparse
import io
import json
import sys
from pathlib import Path

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

GRADE_NAME = {"TS": "특급기밀", "S1": "1급비밀", "S2": "2급대외비", "S3": "3급공개"}


def make_docx(rec: dict, path: Path) -> None:
    from docx import Document
    doc = Document()
    grade = rec.get("target", rec.get("label", "S3"))
    gname = GRADE_NAME.get(grade, grade)
    # 머리말(문서번호·보안등급) — 실제 사내 문서 서식 모사
    h = doc.add_paragraph()
    h.add_run(f"문서번호: {rec.get('doc_id','LK-DOC')}    보안등급: [{gname}]").bold = True
    meta = doc.add_table(rows=1, cols=3)
    meta.style = "Table Grid"
    c = meta.rows[0].cells
    c[0].text = "작성부서: 해당 사업부"; c[1].text = "작성일: 2024-__-__"; c[2].text = f"분류: {gname}({grade})"
    doc.add_heading(rec.get("title", "내부 문서"), level=1)
    for para in str(rec.get("body") or rec.get("text", "")).split("\n"):
        if para.strip():
            doc.add_paragraph(para.strip())
    # 결재란 + 보안 푸터
    appr = doc.add_table(rows=2, cols=3)
    appr.style = "Table Grid"
    appr.rows[0].cells[0].text = "작성"; appr.rows[0].cells[1].text = "검토"; appr.rows[0].cells[2].text = "승인"
    f = doc.add_paragraph()
    f.add_run(f"\n※ 본 문서는 [{gname}]으로 분류됨. 승인된 자 외 열람·유출 금지.").italic = True
    doc.save(str(path))


def make_pdf(rec: dict, path: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    pdfmetrics.registerFont(UnicodeCIDFont("HYSMyeongJo-Medium"))
    grade = rec.get("target", rec.get("label", "S3"))
    gname = GRADE_NAME.get(grade, grade)
    ss = getSampleStyleSheet()
    body_st = ParagraphStyle("kbody", parent=ss["Normal"], fontName="HYSMyeongJo-Medium",
                             fontSize=10, leading=16)
    head_st = ParagraphStyle("khead", parent=ss["Title"], fontName="HYSMyeongJo-Medium", fontSize=15)
    meta_st = ParagraphStyle("kmeta", parent=ss["Normal"], fontName="HYSMyeongJo-Medium", fontSize=9)
    story = [
        Paragraph(f"문서번호: {rec.get('doc_id','LK-DOC')} &nbsp;&nbsp; 보안등급: [{gname}]", meta_st),
        Spacer(1, 4 * mm),
        Paragraph(rec.get("title", "내부 문서"), head_st),
        Spacer(1, 4 * mm),
    ]
    for para in str(rec.get("body") or rec.get("text", "")).split("\n"):
        if para.strip():
            story.append(Paragraph(para.strip().replace("&", "&amp;").replace("<", "&lt;"), body_st))
            story.append(Spacer(1, 2 * mm))
    story.append(Spacer(1, 6 * mm))
    story.append(Paragraph(f"※ 본 문서는 [{gname}]으로 분류됨. 승인된 자 외 열람·유출 금지.", meta_st))
    SimpleDocTemplate(str(path), pagesize=A4, topMargin=18 * mm, bottomMargin=18 * mm).build(story)


def roundtrip_docx(path: Path) -> int:
    """생성한 DOCX를 다시 텍스트로 추출 → M2 파이프라인이 읽을 수 있는지 확인."""
    from docx import Document
    d = Document(str(path))
    return sum(len(p.text) for p in d.paragraphs)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--input", default="datasets/gold/golden50.jsonl")
    ap.add_argument("--limit", type=int, default=0, help="0=전체")
    args = ap.parse_args()

    src = Path(args.input)
    if not src.exists():
        print(f"입력 없음: {src}"); sys.exit(1)
    rows = [json.loads(l) for l in src.read_text(encoding="utf-8").splitlines() if l.strip()]
    if args.limit > 0:
        rows = rows[: args.limit]

    outd = Path("datasets/gold/formats")
    (outd / "docx").mkdir(parents=True, exist_ok=True)
    (outd / "pdf").mkdir(parents=True, exist_ok=True)

    ok_docx = ok_pdf = 0
    for i, rec in enumerate(rows):
        did = rec.get("doc_id", f"DOC-{i:03d}")
        try:
            dp = outd / "docx" / f"{did}.docx"; make_docx(rec, dp)
            rt = roundtrip_docx(dp); ok_docx += 1
        except Exception as exc:  # noqa: BLE001
            print(f"  [{did}] DOCX 오류: {exc}"); rt = -1
        try:
            make_pdf(rec, outd / "pdf" / f"{did}.pdf"); ok_pdf += 1
        except Exception as exc:  # noqa: BLE001
            print(f"  [{did}] PDF 오류: {exc}")
        if i < 3:
            print(f"  [{did}] DOCX 왕복추출 {rt}자 / PDF 생성 | {rec.get('title','')[:30]}")
    print(f"\n생성 완료: DOCX {ok_docx}/{len(rows)} · PDF {ok_pdf}/{len(rows)} → {outd}")
    print("HWP: 한글 독점 포맷 → (a) 한글 SW COM 자동화 (b) HWPX(신형 XML) (c) 발주처 HWP 샘플 — 별도 트랙")


if __name__ == "__main__":
    main()
